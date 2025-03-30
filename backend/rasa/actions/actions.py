import os
import re
from rasa_sdk import Action, Tracker
from rasa_sdk.executor import CollectingDispatcher
from rasa_sdk.events import SlotSet
import fitz  # PyMuPDF
import docx
import csv
from typing import Dict, Text, Any, List

# Constantes para configuración
MAX_RESULTS_PER_PAGE = 3
MAX_MATCHES_PER_DOC = 2
CONTEXT_SIZE = 75
SUPPORTED_FORMATS = {
    '.pdf': 'extract_text_from_pdf',
    '.docx': 'extract_text_from_docx',
    '.txt': 'extract_text_from_txt',
    '.csv': 'extract_text_from_csv'
}

class ActionObtenerDocumento(Action):
    def name(self):
        return "action_obtener_documento"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: dict):
        # Obtener rutas desde variables de entorno o configuración
        docs_dir = os.environ.get("RASA_DOCS_PATH", "documentos/")
        
        # Obtener el nombre del documento desde el slot (si existe)
        doc_name = tracker.get_slot("documento") or "default"
        
        # Limpiar el nombre del documento (eliminar extensión si está incluida)
        doc_name = os.path.splitext(doc_name)[0]
        
        combined_content = ""
        found_files = []
        
        # Verificar si el directorio existe
        if not os.path.exists(docs_dir):
            dispatcher.utter_message(text=f"No se encontró el directorio de documentos: {docs_dir}")
            return []
        
        # Buscar archivos con cualquier extensión soportada
        for ext, extract_method in SUPPORTED_FORMATS.items():
            file_path = os.path.join(docs_dir, f"{doc_name}{ext}")
            if os.path.exists(file_path):
                try:
                    content = getattr(self, extract_method)(file_path)
                    combined_content += content + "\n"
                    found_files.append(os.path.basename(file_path))
                except Exception as e:
                    dispatcher.utter_message(text=f"Error al procesar {os.path.basename(file_path)}: {str(e)}")
        
        if combined_content:
            # Dividir contenido para mostrar solo una parte inicial
            preview = combined_content[:500] + "..." if len(combined_content) > 500 else combined_content
            
            # Guardar el contenido completo en un slot para posibles consultas posteriores
            full_content_event = SlotSet("last_document_content", combined_content)
            file_names_event = SlotSet("last_document_files", found_files)
            
            message = f"Aquí tienes la información sobre '{doc_name}':\n\n{preview}"
            message += "\n\n💡 Si deseas ver más contenido, pídemelo."
            
            dispatcher.utter_message(text=message)
            return [full_content_event, file_names_event]
        else:
            suggestions = self._find_similar_documents(docs_dir, doc_name)
            message = f"No encontré documentos para '{doc_name}'."
            
            if suggestions:
                message += "\n\n¿Quizás quisiste decir alguno de estos?"
                for suggestion in suggestions[:3]:
                    message += f"\n- {suggestion}"
            
            dispatcher.utter_message(text=message)
            return []

    def _find_similar_documents(self, docs_dir, search_name):
        """Encuentra documentos con nombres similares al buscado"""
        all_docs = []
        for ext in SUPPORTED_FORMATS.keys():
            all_docs.extend([os.path.splitext(f)[0] for f in os.listdir(docs_dir) 
                            if f.endswith(ext)])
        
        # Filtrar documentos que contengan parte del nombre buscado
        return [doc for doc in all_docs if search_name.lower() in doc.lower() 
                or doc.lower() in search_name.lower()]

    def extract_text_from_pdf(self, pdf_path):
        try:
            document = fitz.open(pdf_path)
            text = []
            for page_num in range(document.page_count):
                page = document.load_page(page_num)
                text.append(page.get_text())
            return "\n".join(text)
        except Exception as e:
            return f"Error al procesar PDF: {str(e)}"

    def extract_text_from_docx(self, docx_path):
        try:
            document = docx.Document(docx_path)
            text = "\n".join([para.text for para in document.paragraphs])
            return text
        except Exception as e:
            return f"Error al procesar DOCX: {str(e)}"
    
    def extract_text_from_txt(self, txt_path):
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            return f"Error al procesar TXT: {str(e)}"
    
    def extract_text_from_csv(self, csv_path):
        try:
            result = []
            with open(csv_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                headers = next(csv_reader, None)
                if headers:
                    result.append(", ".join(headers))
                for row in csv_reader:
                    result.append(", ".join(row))
            return "\n".join(result)
        except Exception as e:
            return f"Error al procesar CSV: {str(e)}"

class ActionBuscarEnDocumentos(Action):
    def name(self):
        return "action_buscar_en_documentos"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        
        # Obtener la consulta del usuario
        search_term = next((e['value'] for e in tracker.latest_message.get('entities', []) 
                           if e['entity'] == 'search_term'), None)
        
        # Obtener el número de página si estamos paginando resultados
        page = tracker.get_slot("result_page") or 0
        
        # Obtener resultados previos si existen
        previous_results = tracker.get_slot("search_results")
        
        # Si tenemos resultados previos y se solicita "más resultados"
        if not search_term and previous_results and "más" in tracker.latest_message.get('text', '').lower():
            search_term = tracker.get_slot("last_search_term")
            page += 1
            
        if not search_term:
            dispatcher.utter_message(text="¿Qué información específica estás buscando?")
            return []
            
        # Directorio de documentos
        docs_dir = os.environ.get("RASA_DOCS_PATH", "documentos/")
        
        # Si no estamos paginando, buscar nuevos resultados
        if page == 0 or not previous_results:
            # Comprobar si el directorio existe
            if not os.path.exists(docs_dir):
                dispatcher.utter_message(text=f"No se encontró el directorio de documentos: {docs_dir}")
                return []
                
            results = []
            
            # Buscar en todos los formatos soportados
            for file in os.listdir(docs_dir):
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in SUPPORTED_FORMATS:
                    file_path = os.path.join(docs_dir, file)
                    method_name = SUPPORTED_FORMATS[file_ext]
                    matches = getattr(self, f"search_in{file_ext.replace('.', '_')}")(file_path, search_term)
                    if matches:
                        results.append({"file": file, "matches": matches})
        else:
            # Usar resultados previos para paginación
            results = previous_results
        
        # Calcular índices para paginación
        start_idx = page * MAX_RESULTS_PER_PAGE
        end_idx = start_idx + MAX_RESULTS_PER_PAGE
        
        # Responder con los resultados paginados
        if results:
            current_page_results = results[start_idx:end_idx]
            total_pages = (len(results) + MAX_RESULTS_PER_PAGE - 1) // MAX_RESULTS_PER_PAGE
            
            response = f"Resultados para '{search_term}' (página {page+1} de {total_pages}):\n\n"
            
            for result in current_page_results:
                response += f"📄 **{os.path.splitext(result['file'])[0]}**:\n"
                for i, match in enumerate(result['matches'][:MAX_MATCHES_PER_DOC]):
                    response += f"- {match}\n"
                response += "\n"
                
            # Agregar indicador de paginación si hay más páginas
            if page + 1 < total_pages:
                response += "💡 Puedes pedirme \"más resultados\" para ver la siguiente página."
                
            # Guardar estado para paginación
            dispatcher.utter_message(text=response)
            return [
                SlotSet("search_results", results),
                SlotSet("result_page", page),
                SlotSet("last_search_term", search_term)
            ]
        else:
            dispatcher.utter_message(text=f"No encontré información sobre '{search_term}' en los documentos disponibles.")
            return [SlotSet("result_page", 0)]
    
    def search_in_pdf(self, pdf_path, search_term):
        matches = []
        try:
            document = fitz.open(pdf_path)
            for page_num in range(document.page_count):
                page = document.load_page(page_num)
                text = page.get_text()
                paragraphs = text.split('\n\n')
                
                for para in paragraphs:
                    if search_term.lower() in para.lower():
                        # Extraer contexto alrededor del término
                        matches.append(self.extract_context(para, search_term))
                        
            return matches[:5]  # Limitar a 5 coincidencias
        except Exception as e:
            matches.append(f"Error al procesar PDF: {str(e)}")
            return matches
    
    def search_in_docx(self, docx_path, search_term):
        matches = []
        try:
            document = docx.Document(docx_path)
            for para in document.paragraphs:
                if search_term.lower() in para.text.lower():
                    matches.append(self.extract_context(para.text, search_term))
            return matches[:5]  # Limitar a 5 coincidencias
        except Exception as e:
            matches.append(f"Error al procesar DOCX: {str(e)}")
            return matches
    
    def search_in_txt(self, txt_path, search_term):
        matches = []
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    if search_term.lower() in para.lower():
                        matches.append(self.extract_context(para, search_term))
            return matches[:5]
        except Exception as e:
            matches.append(f"Error al procesar TXT: {str(e)}")
            return matches
            
    def search_in_csv(self, csv_path, search_term):
        matches = []
        try:
            with open(csv_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                headers = next(csv_reader, None)
                
                for row_num, row in enumerate(csv_reader, 1):
                    row_text = ", ".join(row)
                    if search_term.lower() in row_text.lower():
                        if headers:
                            matches.append(f"Fila {row_num}: {row_text}")
                        else:
                            matches.append(row_text)
            return matches[:5]
        except Exception as e:
            matches.append(f"Error al procesar CSV: {str(e)}")
            return matches
    
    def extract_context(self, text, term):
        """Extrae un contexto mejorado alrededor del término de búsqueda"""
        # Encontrar todas las ocurrencias del término (ignorando mayúsculas/minúsculas)
        pattern = re.compile(re.escape(term), re.IGNORECASE)
        matches = list(pattern.finditer(text))
        
        if not matches:
            return text[:150] + "..." if len(text) > 150 else text
        
        # Extracción de contexto inteligente (párrafo completo o limitar por tamaño)
        paragraphs = text.split('\n')
        for para in paragraphs:
            if term.lower() in para.lower():
                if len(para) <= CONTEXT_SIZE * 3:
                    return para  # Devolver párrafo completo si es pequeño
                
                # Si el párrafo es largo, extraer contexto alrededor del primer match
                index = matches[0].start()
                start = max(0, index - CONTEXT_SIZE)
                end = min(len(text), index + len(term) + CONTEXT_SIZE)
                
                result = ""
                if start > 0:
                    result += "..."
                result += text[start:end]
                if end < len(text):
                    result += "..."
                return result
                
        # Si no encontramos el término en los párrafos (raro, pero por si acaso)
        index = matches[0].start()
        start = max(0, index - CONTEXT_SIZE)
        end = min(len(text), index + len(term) + CONTEXT_SIZE)
        
        result = ""
        if start > 0:
            result += "..."
        result += text[start:end]
        if end < len(text):
            result += "..."
            
        return result

class ActionListarDocumentos(Action):
    def name(self):
        return "action_listar_documentos"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: dict):
        # Obtener ruta desde variables de entorno o configuración
        docs_dir = os.environ.get("RASA_DOCS_PATH", "documentos/")
        
        # Verificar si el directorio existe
        if not os.path.exists(docs_dir):
            dispatcher.utter_message(text=f"No se encontró el directorio de documentos: {docs_dir}")
            return []
        
        # Listar archivos de todos los formatos soportados
        doc_files = {}
        for ext in SUPPORTED_FORMATS.keys():
            files = [f for f in os.listdir(docs_dir) if f.lower().endswith(ext)]
            if files:
                doc_files[ext] = files
        
        if not doc_files:
            dispatcher.utter_message(text="No hay documentos disponibles.")
            return []
        
        # Construir mensaje de respuesta
        response = "Estos son los documentos disponibles:\n\n"
        
        # Iconos para cada formato
        format_icons = {
            '.pdf': '📑',
            '.docx': '📝',
            '.txt': '📄',
            '.csv': '📊'
        }
        
        # Mostrar documentos por formato
        for ext, files in doc_files.items():
            if files:
                icon = format_icons.get(ext, '📁')
                ext_name = ext.upper()[1:]  # Quitar el punto y convertir a mayúsculas
                response += f"{icon} **Documentos {ext_name}**:\n"
                for i, file in enumerate(sorted(files), 1):
                    # Eliminar la extensión para obtener el nombre base
                    name = os.path.splitext(file)[0]
                    response += f"{i}. {name}\n"
                response += "\n"
        
        response += "💡 Puedes pedirme información sobre cualquiera de estos documentos o buscar un tema específico."
        dispatcher.utter_message(text=response)
        return []

class ActionMasContenido(Action):
    def name(self):
        return "action_mostrar_mas_contenido"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: dict):
        # Recuperar el contenido completo del documento
        content = tracker.get_slot("last_document_content")
        file_names = tracker.get_slot("last_document_files") or []
        
        if not content:
            dispatcher.utter_message(text="No hay ningún documento cargado actualmente. ¿Quieres consultar algún documento específico?")
            return []
        
        # Determinar la parte a mostrar (siguiente segmento)
        current_offset = tracker.get_slot("content_offset") or 500
        next_chunk = content[current_offset:current_offset+1000]
        
        if next_chunk:
            files_text = ", ".join(file_names) if file_names else "documento"
            message = f"Continuando con el contenido del {files_text}:\n\n{next_chunk}"
            
            # Si hay más contenido por mostrar
            if current_offset + 1000 < len(content):
                message += "\n\n💡 Puedes seguir pidiendo \"más contenido\" para continuar leyendo."
                
            dispatcher.utter_message(text=message)
            return [SlotSet("content_offset", current_offset + 1000)]
        else:
            dispatcher.utter_message(text="Has llegado al final del documento.")
            return [SlotSet("content_offset", 0)]  # Reiniciar para futuros documentos